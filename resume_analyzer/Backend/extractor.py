"""
extractor.py
============
Responsible for ONE thing: taking a file path (PDF or DOCX)
and returning clean raw text.

WHY THIS IS ITS OWN MODULE:
    Separation of concerns. The NLP pipeline should never know
    or care about file formats. It only speaks "text".
"""

import re
import pdfplumber
from docx import Document # type: ignore
from pathlib import Path


# MAIN PUBLIC FUNCTION — this is what every other module calls
def extract_text(file_path: str) -> dict: # type: ignore
    """
    Entry point. Detects format and routes to the right extractor.

    Returns a dict so we carry metadata alongside the text:
    {
        "raw_text": "John Doe\nSoftware Engineer...",
        "file_type": "pdf",
        "page_count": 2,          # PDF only
        "char_count": 3241,
        "error": None             # or an error message string
    }

    WHY RETURN A DICT INSTEAD OF JUST A STRING?
        The page_count and char_count become signals for ATS scoring
        later. A 5-page resume is a red flag; a 200-char resume
        probably failed to extract. Better to carry that info now.
    """
    path = Path(file_path)

    if not path.exists():
        return _error_result(f"File not found: {file_path}") # type: ignore

    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return _extract_pdf(path) # type: ignore
    elif suffix in (".docx", ".doc"):
        return _extract_docx(path) # type: ignore
    elif suffix == ".txt":
        return _extract_txt(path) # type: ignore
    else:
        return _error_result(f"Unsupported file type: {suffix}") # type: ignore


# PDF EXTRACTION
def _extract_pdf(path: Path) -> dict: # type: ignore
    """
    pdfplumber reads PDF pages and extracts text with layout awareness.

    WHY pdfplumber OVER PyPDF2?
        pdfplumber uses pdfminer under the hood which understands
        text position (x, y coordinates). This means it can
        reconstruct columns, tables, and headers correctly.
        PyPDF2 often mangles multi-column resumes.

    WHY extract_text(x_tolerance, y_tolerance)?
        x_tolerance: characters within this many points horizontally
                     are considered on the same "word"
        y_tolerance: lines within this many points vertically
                     are considered the same "line"
        Without these, you get random spaces mid-word from PDF kerning.
    """
    pages_text = []

    try:
        with pdfplumber.open(path) as pdf:
            page_count = len(pdf.pages)

            for page in pdf.pages:
                # extract_text returns None if a page has no text
                # (e.g. a scanned image — we handle that gracefully)
                text = page.extract_text(x_tolerance=3, y_tolerance=3)
                if text:
                    pages_text.append(text) # type: ignore

        raw_text = "\n".join(pages_text) # type: ignore
        cleaned = _clean_text(raw_text)

        return {
            "raw_text": cleaned,
            "file_type": "pdf",
            "page_count": page_count,
            "char_count": len(cleaned),
            "error": None,
        } # type: ignore

    except Exception as e:
        return _error_result(f"PDF extraction failed: {str(e)}") # type: ignore

# DOCX EXTRACTION
def _extract_docx(path: Path) -> dict: # type: ignore
    """
    python-docx gives us a Document object with:
        - doc.paragraphs  → list of Paragraph objects (main body text)
        - doc.tables      → list of Table objects (skills tables, etc.)
        - paragraph.runs  → the styled chunks within a paragraph

    WHY DO WE ALSO READ TABLES?
        Many resumes put skills in a table grid:
        | Python | Django | PostgreSQL |
        | React  | Docker | AWS        |
        If we only read paragraphs, those skills are invisible to us.

    WHY READ paragraph.style.name?
        DOCX styles (Heading 1, Heading 2, Normal) are section
        markers. A paragraph styled "Heading 1" is almost certainly
        a section header like "EXPERIENCE" or "EDUCATION".
        We'll use this in section_detector.py later.
    """
    try:
        doc = Document(path) # type: ignore
        parts = []

        # 1. Read all paragraphs
        for para in doc.paragraphs: # type: ignore
            text = para.text.strip() # type: ignore
            if text:
                parts.append(text) # type: ignore

        # 2. Read all tables (captures skill grids, etc.)
        for table in doc.tables: # type: ignore
            for row in table.rows: # type: ignore
                row_text = " | ".join(
                    cell.text.strip() # type: ignore
                    for cell in row.cells # type: ignore
                    if cell.text.strip() # type: ignore
                )
                if row_text:
                    parts.append(row_text) # type: ignore

        raw_text = "\n".join(parts) # type: ignore
        cleaned = _clean_text(raw_text)

        return {
            "raw_text": cleaned,
            "file_type": "docx",
            "page_count": None,   # DOCX doesn't have a reliable page count
            "char_count": len(cleaned),
            "error": None,
        } # type: ignore

    except Exception as e:
        return _error_result(f"DOCX extraction failed: {str(e)}") # type: ignore



# TXT EXTRACTION
def _extract_txt(path: Path) -> dict: # type: ignore
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            raw_text = f.read()
        
        cleaned = _clean_text(raw_text)
        return {
            "raw_text": cleaned,
            "file_type": "txt",
            "page_count": None,
            "char_count": len(cleaned),
            "error": None,
        }
    except Exception as e:
        return _error_result(f"TXT extraction failed: {str(e)}") # type: ignore


# TEXT CLEANING
def _clean_text(text: str) -> str:
    """
    Normalises raw extracted text so NLP tools work consistently.

    WHAT WE DO AND WHY:

    1. Unicode normalisation
       PDFs often contain special Unicode dashes (—, –), bullets (•·▪),
       and smart quotes (""). spaCy handles these, but consistent
       ASCII simplifies regex patterns in section_detector.py.

    2. Remove non-printable characters
       PDF extraction sometimes pulls in null bytes, form-feed
       characters, and other control characters. They're invisible
       but break string operations.

    3. Collapse multiple blank lines
       Resume extractors tend to output 4-5 empty lines between
       sections. We normalise to max 2, which preserves structure
       without bloating the text.

    4. Strip leading/trailing whitespace per line
       Hanging spaces confuse section-header regex patterns.

    WHAT WE DON'T DO:
       We don't lowercase here. Section detection and NER both
       need case information ("Python" vs "python" matters for NER;
       "EXPERIENCE" all-caps is a strong section signal).
    """
    if not text:
        return ""

    # Replace special Unicode dashes and bullets with ASCII equivalents
    text = text.replace("\u2013", "-").replace("\u2014", "-")
    text = text.replace("\u2022", "-").replace("\u00b7", "-")
    text = text.replace("\u2019", "'").replace("\u201c", '"').replace("\u201d", '"')

    # Remove non-printable / control characters (keep newlines and tabs)
    text = re.sub(r"[^\x09\x0A\x0D\x20-\x7E]", " ", text)

    # Strip each line, remove lines that are pure whitespace
    lines = [line.strip() for line in text.splitlines()]
    lines = [line for line in lines if line]

    # Collapse sequences of 3+ blank lines into 2
    # (We reinsert blank lines at section breaks, not here)
    result_lines = []
    blank_count = 0
    for line in lines:
        if line == "":
            blank_count += 1
            if blank_count <= 2:
                result_lines.append(line) # type: ignore
        else:
            blank_count = 0
            result_lines.append(line) # type: ignore

    return "\n".join(result_lines).strip() # type: ignore


# HELPERS
def _error_result(message: str) -> dict: # type: ignore
    return {
        "raw_text": "",
        "file_type": "unknown",
        "page_count": None,
        "char_count": 0,
        "error": message,
    } # type: ignore

if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage: python extractor.py <path_to_resume.pdf_or_docx>")
        sys.exit(1)

    result = extract_text(sys.argv[1]) # type: ignore

    print("=== EXTRACTION RESULT ===")
    print(f"File type  : {result['file_type']}")
    print(f"Page count : {result['page_count']}")
    print(f"Char count : {result['char_count']}")
    print(f"Error      : {result['error']}")
    print("\n=== FIRST 500 CHARS ===")
    print(result["raw_text"][:500]) # type: ignore